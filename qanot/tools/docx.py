"""Word (.docx) document tools — create, read, edit."""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path

from qanot.registry import ToolRegistry

logger = logging.getLogger(__name__)


def _add_formatted_text(paragraph, text: str) -> None:
    """Parse simple bold (**text**) formatting."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def register_docx_tools(registry: ToolRegistry, workspace_dir: str) -> None:
    """Register Word document tools."""

    # ── create_docx ──
    async def create_docx(params: dict) -> str:
        """Create a Word document."""
        try:
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            return json.dumps({"error": "python-docx kutubxonasi o'rnatilmagan. pip install python-docx"})

        filename = params.get("filename", "document.docx")
        if not filename.endswith(".docx"):
            filename += ".docx"

        title = params.get("title", "")
        content = params.get("content", "")
        rows = params.get("rows")

        doc = Document()

        # Title
        if title:
            p = doc.add_heading(title, level=0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Content — parse simple markdown-like format
        if content:
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif re.match(r"^\d+\.\s", line):
                    doc.add_paragraph(re.sub(r"^\d+\.\s*", "", line), style="List Number")
                else:
                    p = doc.add_paragraph()
                    _add_formatted_text(p, line)

        # Table
        if rows and isinstance(rows, list) and len(rows) > 0:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    table.rows[i].cells[j].text = str(cell_text)
                    if i == 0:
                        for paragraph in table.rows[i].cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        # Save
        filepath = Path(workspace_dir) / filename
        doc.save(str(filepath))

        return json.dumps({
            "status": "ok",
            "file": str(filepath),
            "filename": filename,
            "message": f"{filename} yaratildi",
        })

    # ── read_docx ──
    async def read_docx(params: dict) -> str:
        """Read content from a Word document."""
        try:
            from docx import Document
        except ImportError:
            return json.dumps({"error": "python-docx kutubxonasi o'rnatilmagan"})

        filepath = params.get("file", "")
        if not filepath:
            return json.dumps({"error": "file parametri kerak"})

        path = Path(filepath) if Path(filepath).is_absolute() else Path(workspace_dir) / filepath
        if not path.exists():
            return json.dumps({"error": f"Fayl topilmadi: {filepath}"})

        doc = Document(str(path))
        content = []
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                content.append(f"{'#' * level} {para.text}")
            elif para.style.name == "List Bullet":
                content.append(f"- {para.text}")
            elif para.style.name == "List Number":
                content.append(f"1. {para.text}")
            else:
                content.append(para.text)

        # Read tables
        tables_text = []
        for i, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                table_rows.append([cell.text for cell in row.cells])
            tables_text.append({"table_index": i, "rows": table_rows})

        return json.dumps({
            "content": "\n".join(content),
            "tables": tables_text,
            "paragraphs": len(doc.paragraphs),
        }, ensure_ascii=False)

    # ── edit_docx ──
    async def edit_docx(params: dict) -> str:
        """Edit an existing Word document — append text, replace text, or add table."""
        try:
            from docx import Document
            from docx.shared import Pt  # noqa: F841
        except ImportError:
            return json.dumps({"error": "python-docx kutubxonasi o'rnatilmagan"})

        filepath = params.get("file", "")
        if not filepath:
            return json.dumps({"error": "file parametri kerak"})

        path = Path(filepath) if Path(filepath).is_absolute() else Path(workspace_dir) / filepath
        if not path.exists():
            return json.dumps({"error": f"Fayl topilmadi: {filepath}"})

        doc = Document(str(path))
        action = params.get("action", "append")  # append, replace, add_table

        if action == "append":
            content = params.get("content", "")
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    p = doc.add_paragraph()
                    _add_formatted_text(p, line)

        elif action == "replace":
            old_text = params.get("old_text", "")
            new_text = params.get("new_text", "")
            if not old_text:
                return json.dumps({"error": "old_text kerak"})
            count = 0
            for para in doc.paragraphs:
                if old_text in para.text:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            count += 1
            if count == 0:
                return json.dumps({"error": f"'{old_text}' topilmadi"})

        elif action == "add_table":
            rows = params.get("rows", [])
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    for j, val in enumerate(row):
                        table.rows[i].cells[j].text = str(val)

        doc.save(str(path))
        return json.dumps({"status": "ok", "message": f"{path.name} yangilandi", "action": action})

    # Register tools
    registry.register(
        name="create_docx",
        description=(
            "Word (.docx) hujjat yaratish. Shartnoma, hisobot, taklifnoma va boshqa hujjatlar uchun. "
            "Markdown formatda content yozing: # sarlavha, ## kichik sarlavha, **qalin**, - ro'yxat. "
            "Jadval uchun rows parametrini ishlating."
        ),
        parameters={
            "type": "object",
            "required": ["filename"],
            "properties": {
                "filename": {"type": "string", "description": "Fayl nomi (masalan: shartnoma.docx)"},
                "title": {"type": "string", "description": "Hujjat sarlavhasi"},
                "content": {"type": "string", "description": "Hujjat matni (Markdown format)"},
                "rows": {
                    "type": "array",
                    "description": "Jadval ma'lumotlari. Birinchi qator — sarlavha. [[\"Nomi\",\"Narxi\"],[\"Tovar1\",\"5000\"]]",
                    "items": {"type": "array", "items": {"type": "string"}},
                },
            },
        },
        handler=create_docx,
    )

    registry.register(
        name="read_docx",
        description="Word (.docx) hujjatni o'qish. Matn, sarlavhalar, jadvallarni qaytaradi.",
        parameters={
            "type": "object",
            "required": ["file"],
            "properties": {
                "file": {"type": "string", "description": "Fayl nomi yoki to'liq path"},
            },
        },
        handler=read_docx,
    )

    registry.register(
        name="edit_docx",
        description=(
            "Word (.docx) hujjatni tahrirlash. Matn qo'shish (append), "
            "almashtirish (replace), jadval qo'shish (add_table)."
        ),
        parameters={
            "type": "object",
            "required": ["file"],
            "properties": {
                "file": {"type": "string", "description": "Fayl nomi yoki to'liq path"},
                "action": {
                    "type": "string",
                    "enum": ["append", "replace", "add_table"],
                    "description": "Harakat turi: append (matn qo'shish), replace (almashtirish), add_table (jadval)",
                },
                "content": {"type": "string", "description": "Qo'shiladigan matn (append uchun, Markdown format)"},
                "old_text": {"type": "string", "description": "Almashtirilishi kerak bo'lgan matn (replace uchun)"},
                "new_text": {"type": "string", "description": "Yangi matn (replace uchun)"},
                "rows": {
                    "type": "array",
                    "description": "Jadval qatorlari (add_table uchun)",
                    "items": {"type": "array", "items": {"type": "string"}},
                },
            },
        },
        handler=edit_docx,
    )

    logger.info("DOCX tools registered: create_docx, read_docx, edit_docx")
